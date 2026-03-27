"""Document parsing — extracts text from various file formats."""

import logging
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


@dataclass
class DocumentSection:
    text: str
    page_number: int | None = None


@dataclass
class LoadedDocument:
    sections: list[DocumentSection]

    @property
    def text(self) -> str:
        return "\n\n".join(section.text for section in self.sections if section.text.strip())


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Unable to decode text document")


async def load_document(filename: str, content: bytes) -> LoadedDocument:
    """Extract text content from an uploaded document."""
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    logger.info("Loading document", extra={"document_filename": filename, "type": ext})

    if ext in {".txt", ".md"}:
        return LoadedDocument(sections=[DocumentSection(text=_decode_text(content))])

    if ext == ".pdf":
        reader = PdfReader(BytesIO(content))
        sections = [
            DocumentSection(text=(page.extract_text() or "").strip(), page_number=index + 1)
            for index, page in enumerate(reader.pages)
        ]
        return LoadedDocument(sections=[section for section in sections if section.text])

    if ext == ".docx":
        document = DocxDocument(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        return LoadedDocument(sections=[DocumentSection(text=text)])

    raise ValueError(f"Unsupported file type: {ext}")
